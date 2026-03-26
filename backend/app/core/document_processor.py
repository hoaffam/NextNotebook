"""
Document Processor
Extract text from PDF, DOCX, and TXT files with structured metadata
"""

import os
from typing import Optional, Dict, List, Any, Union
from app.utils.logger import setup_logger

logger = setup_logger(__name__)


class DocumentProcessor:
    """Process and extract text from various document formats"""

    def __init__(self):
        self.supported_formats = ["pdf", "docx", "txt"]

    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a document file (legacy method - backward compatible)

        Args:
            file_path: Path to the document file

        Returns:
            Extracted text content
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = file_path.split(".")[-1].lower()

        if ext == "pdf":
            return self._extract_pdf(file_path)
        elif ext == "docx":
            return self._extract_docx(file_path)
        elif ext == "txt":
            return self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def extract_structured(self, file_path: str) -> Dict[str, Any]:
        """
        Extract structured data with metadata for smart chunking

        Args:
            file_path: Path to the document file

        Returns:
            Dictionary with:
                - file_type: str (pdf, docx, txt)
                - data: Structured data specific to file type
                - plain_text: str (for backward compatibility)
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = file_path.split(".")[-1].lower()

        if ext == "pdf":
            pages_data = self._extract_pdf_structured(file_path)
            plain_text = "\n\n".join([p["text"] for p in pages_data])
            return {
                "file_type": "pdf",
                "data": pages_data,
                "plain_text": plain_text
            }
        elif ext == "docx":
            sections_data = self._extract_docx_structured(file_path)
            plain_text = self._flatten_docx_sections(sections_data)
            return {
                "file_type": "docx",
                "data": sections_data,
                "plain_text": plain_text
            }
        elif ext == "txt":
            text = self._extract_txt(file_path)
            return {
                "file_type": "txt",
                "data": text,  # TXT is already simple
                "plain_text": text
            }
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF using pdfplumber"""
        try:
            import pdfplumber
            
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
            
            return "\n\n".join(text_parts)
            
        except ImportError:
            logger.warning("pdfplumber not installed, trying PyPDF2")
            return self._extract_pdf_fallback(file_path)
    
    def _extract_pdf_fallback(self, file_path: str) -> str:
        """Fallback PDF extraction using PyPDF2"""
        from PyPDF2 import PdfReader
        
        reader = PdfReader(file_path)
        text_parts = []
        
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
        
        return "\n\n".join(text_parts)
    
    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        from docx import Document
        
        doc = Document(file_path)
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)
        
        return "\n\n".join(paragraphs)
    
    def _extract_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
        
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Could not decode text file with any supported encoding")
    
    def get_file_info(self, file_path: str) -> dict:
        """Get information about a file"""
        stat = os.stat(file_path)
        ext = file_path.split(".")[-1].lower()

        return {
            "path": file_path,
            "filename": os.path.basename(file_path),
            "extension": ext,
            "size_bytes": stat.st_size,
            "size_mb": round(stat.st_size / (1024 * 1024), 2)
        }

    # ============================================================
    # STRUCTURED EXTRACTION METHODS
    # ============================================================

    def _extract_pdf_structured(self, file_path: str) -> List[Dict]:
        """
        Extract PDF with page metadata

        Returns:
            List of page dictionaries:
            {
                "page_number": int,
                "text": str,
                "char_start": int,  # Global position
                "char_end": int
            }
        """
        try:
            import pdfplumber

            pages_data = []
            cumulative_chars = 0

            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        # Remove the [Page X] marker for cleaner text
                        char_start = cumulative_chars
                        char_end = cumulative_chars + len(page_text)

                        pages_data.append({
                            "page_number": page_num + 1,
                            "text": page_text,
                            "char_start": char_start,
                            "char_end": char_end
                        })

                        cumulative_chars = char_end + 2  # +2 for \n\n separator

            logger.info(f"Extracted {len(pages_data)} pages with metadata from PDF")
            return pages_data

        except ImportError:
            logger.warning("pdfplumber not installed, using fallback")
            return self._extract_pdf_structured_fallback(file_path)

    def _extract_pdf_structured_fallback(self, file_path: str) -> List[Dict]:
        """Fallback PDF extraction using PyPDF2"""
        from PyPDF2 import PdfReader

        reader = PdfReader(file_path)
        pages_data = []
        cumulative_chars = 0

        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                char_start = cumulative_chars
                char_end = cumulative_chars + len(page_text)

                pages_data.append({
                    "page_number": page_num + 1,
                    "text": page_text,
                    "char_start": char_start,
                    "char_end": char_end
                })

                cumulative_chars = char_end + 2

        return pages_data

    def _extract_docx_structured(self, file_path: str) -> List[Dict]:
        """
        Extract DOCX with heading hierarchy

        Returns:
            List of section dictionaries:
            {
                "heading": str,
                "heading_level": int,
                "section_path": str,  # Breadcrumb
                "paragraphs": List[str],
                "char_start": int,
                "char_end": int,
                "tables": List[Dict]
            }
        """
        from docx import Document

        doc = Document(file_path)
        sections = []
        current_section = None
        heading_stack = []  # Track heading hierarchy
        cumulative_chars = 0

        for element in doc.element.body:
            # Check if paragraph
            if element.tag.endswith('p'):
                para = None
                for p in doc.paragraphs:
                    if p._element == element:
                        para = p
                        break

                if not para or not para.text.strip():
                    continue

                # Check if it's a heading
                if para.style.name.startswith("Heading"):
                    # Extract heading level
                    try:
                        level = int(para.style.name.split()[-1])
                    except:
                        level = 1

                    # Update heading stack (maintain hierarchy)
                    heading_stack = [h for h in heading_stack if h["level"] < level]
                    heading_stack.append({"level": level, "text": para.text})

                    # Create section path (breadcrumb)
                    section_path = " > ".join([h["text"] for h in heading_stack])

                    # Save previous section if exists
                    if current_section and current_section.get("paragraphs"):
                        current_section["char_end"] = cumulative_chars
                        sections.append(current_section)

                    # Start new section
                    current_section = {
                        "heading": para.text,
                        "heading_level": level,
                        "section_path": section_path,
                        "paragraphs": [],
                        "tables": [],
                        "char_start": cumulative_chars,
                        "char_end": cumulative_chars  # Will be updated
                    }
                else:
                    # Regular paragraph
                    if current_section is None:
                        # No heading yet, create default section
                        current_section = {
                            "heading": "(No Heading)",
                            "heading_level": 0,
                            "section_path": "(No Heading)",
                            "paragraphs": [],
                            "tables": [],
                            "char_start": cumulative_chars,
                            "char_end": cumulative_chars
                        }

                    current_section["paragraphs"].append(para.text)
                    cumulative_chars += len(para.text) + 2  # +2 for \n\n

            # Check if table
            elif element.tag.endswith('tbl'):
                table_obj = None
                for tbl in doc.tables:
                    if tbl._element == element:
                        table_obj = tbl
                        break

                if table_obj:
                    # Extract table as text
                    table_text_rows = []
                    for row in table_obj.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        if row_text.strip():
                            table_text_rows.append(row_text)

                    table_text = "\n".join(table_text_rows)

                    if current_section is None:
                        current_section = {
                            "heading": "(No Heading)",
                            "heading_level": 0,
                            "section_path": "(No Heading)",
                            "paragraphs": [],
                            "tables": [],
                            "char_start": cumulative_chars,
                            "char_end": cumulative_chars
                        }

                    current_section["tables"].append({
                        "name": f"Table {len(current_section['tables']) + 1}",
                        "text": table_text,
                        "char_start": cumulative_chars,
                        "char_end": cumulative_chars + len(table_text)
                    })

                    cumulative_chars += len(table_text) + 2

        # Don't forget last section
        if current_section and (current_section.get("paragraphs") or current_section.get("tables")):
            current_section["char_end"] = cumulative_chars
            sections.append(current_section)

        # If no sections found (no headings), create one section with all paragraphs
        if not sections:
            all_paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            if all_paragraphs:
                sections.append({
                    "heading": "(Document)",
                    "heading_level": 0,
                    "section_path": "(Document)",
                    "paragraphs": all_paragraphs,
                    "tables": [],
                    "char_start": 0,
                    "char_end": sum(len(p) for p in all_paragraphs)
                })

        logger.info(f"Extracted {len(sections)} sections with headings from DOCX")
        return sections

    def _flatten_docx_sections(self, sections_data: List[Dict]) -> str:
        """Convert DOCX sections back to plain text"""
        parts = []

        for section in sections_data:
            # Add heading
            heading = section.get("heading", "")
            if heading and heading != "(No Heading)" and heading != "(Document)":
                parts.append(f"# {heading}")

            # Add paragraphs
            paragraphs = section.get("paragraphs", [])
            if paragraphs:
                parts.append("\n\n".join(paragraphs))

            # Add tables
            tables = section.get("tables", [])
            for table in tables:
                parts.append(f"\n{table.get('name', 'Table')}:\n{table.get('text', '')}")

        return "\n\n".join(parts)
